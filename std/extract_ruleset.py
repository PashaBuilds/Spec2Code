"""Coding-standard Word doc -> machine-readable ruleset (Brief §16).

Pipeline: python-docx extracts text + tables -> (if LLM enabled) structure into a ruleset
-> human approval. Without an LLM, this returns the default ruleset with the extracted raw
text attached under `_source_text` for a human to turn into rules.

Usage:
    python -m std.extract_ruleset path/to/standard.docx > std/extracted.ruleset.json
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parent.parent
_DEFAULT = _ROOT / "std" / "default.ruleset.json"


def extract_text(docx_path: Path) -> str:
    """Pull paragraphs and table cells out of a .docx using python-docx."""
    from docx import Document  # lazy import

    doc = Document(str(docx_path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def build_ruleset(docx_path: Path, *, spec_llm: dict | None = None) -> dict:
    """Return a ruleset dict from the Word document.

    With an LLM endpoint configured, structures the extracted text into the ruleset schema.
    Otherwise returns the default ruleset annotated with the raw text for human review.
    """
    text = extract_text(docx_path)
    base = json.loads(_DEFAULT.read_text(encoding="utf-8"))

    if spec_llm and spec_llm.get("enabled"):
        from orchestrator.llm.client import LlmClient, LlmConfig

        client = LlmClient(LlmConfig.resolve(spec_llm))
        if client.available:
            messages = [
                {"role": "system", "content":
                 "Convert this embedded-C coding standard into a JSON ruleset matching the keys: "
                 "formatting, naming (hungarian_prefixes, function_pattern, function_regex), "
                 "print_conventions, doxygen. Return ONLY JSON."},
                {"role": "user", "content": text},
            ]
            try:
                out = client.chat(messages, temperature=0.0)
                base = json.loads(out[out.find("{"): out.rfind("}") + 1])
            except Exception:
                base["_source_text"] = text  # fall back; keep text for human review
                return base
            base["_needs_human_review"] = True
            return base

    base["_source_text"] = text
    base["_needs_human_review"] = True
    return base


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("usage: python -m std.extract_ruleset <standard.docx>", file=sys.stderr)
        raise SystemExit(2)
    print(json.dumps(build_ruleset(Path(sys.argv[1])), indent=2))
