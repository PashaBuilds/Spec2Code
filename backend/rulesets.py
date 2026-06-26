"""Coding-standard ruleset import, validation, and persistence.

The LLM is intentionally only an extractor. The canonical contract is the
versioned ruleset schema plus deterministic validation/sample checks.
"""

from __future__ import annotations

import base64
import copy
import io
import json
import os
import re
import sys
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from jsonschema import Draft7Validator

from orchestrator.llm.client import LlmClient, LlmConfig, LlmError
from orchestrator.qc import naming_linter

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_RULESET_PATH = ROOT / "std" / "default.ruleset.json"
SCHEMA_PATH = ROOT / "schemas" / "ruleset.schema.json"

RULESET_SCHEMA = json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))
DEFAULT_RULESET = json.loads(DEFAULT_RULESET_PATH.read_text(encoding="utf-8"))
_VALIDATOR = Draft7Validator(RULESET_SCHEMA)

_ALLOWED_TOP = {"ruleset_version", "formatting", "naming", "print_conventions", "doxygen"}
_ALLOWED_NESTED = {
    "formatting": {"brace_style", "indent", "line_ending", "max_line_length"},
    "naming": {"hungarian_prefixes", "function_pattern", "function_regex"},
    "print_conventions": {"line_terminator_in_prints"},
    "doxygen": {"required_on_public_functions"},
}


def _issue(path: str, message: str, severity: str = "error") -> dict[str, str]:
    return {"severity": severity, "path": path, "message": message}


def _coerce_value(path: str, value: Any) -> Any:
    if path == "formatting/max_line_length" and isinstance(value, str):
        try:
            return int(value.strip())
        except ValueError:
            return value
    if path == "doxygen/required_on_public_functions" and isinstance(value, str):
        lowered = value.strip().lower()
        if lowered in {"true", "yes", "1", "required", "mandatory"}:
            return True
        if lowered in {"false", "no", "0", "optional", "not_required"}:
            return False
    if path == "print_conventions/line_terminator_in_prints" and value == "\r\n":
        return "\\r\\n"
    return value


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1254", "windows-1254", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_docx_text(raw: bytes) -> str:
    doc = Document(io.BytesIO(raw))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def source_text_from_upload(filename: str, *, text: str = "", content_base64: str = "") -> str:
    name = filename.lower()
    if content_base64:
        try:
            raw = base64.b64decode(content_base64, validate=True)
        except ValueError as exc:
            raise ValueError("file content is not valid base64") from exc
        if name.endswith(".docx"):
            return _extract_docx_text(raw)
        return _decode_text(raw)
    return text


def _json_object_from_text(text: str) -> dict:
    stripped = text.strip()
    if stripped.startswith("```"):
        stripped = re.sub(r"^\s*```(?:json)?\s*", "", stripped, flags=re.IGNORECASE)
        stripped = re.sub(r"\s*```\s*$", "", stripped)
    start = stripped.find("{")
    end = stripped.rfind("}")
    if start < 0 or end < start:
        raise ValueError("LLM response did not contain a JSON object")
    data = json.loads(stripped[start:end + 1])
    if not isinstance(data, dict):
        raise ValueError("ruleset JSON root must be an object")
    return data


def _merge_known(default: dict, candidate: dict, path: str = "") -> tuple[dict, list[dict[str, str]]]:
    merged = copy.deepcopy(default)
    issues: list[dict[str, str]] = []
    for key, value in candidate.items():
        key_path = f"{path}/{key}".strip("/")
        if path == "" and key not in _ALLOWED_TOP:
            issues.append(_issue(key_path, "ignored unsupported top-level ruleset field", "warning"))
            continue
        if key in _ALLOWED_NESTED and isinstance(value, dict):
            for sub_key, sub_value in value.items():
                sub_path = f"{key}/{sub_key}"
                if sub_key not in _ALLOWED_NESTED[key]:
                    issues.append(_issue(sub_path, "ignored unsupported ruleset field", "warning"))
                    continue
                if key == "naming" and sub_key == "hungarian_prefixes" and isinstance(sub_value, dict):
                    merged[key][sub_key] = {str(k): str(v) for k, v in sub_value.items()}
                else:
                    merged[key][sub_key] = _coerce_value(sub_path, sub_value)
            continue
        if key in merged:
            merged[key] = value
    return merged, issues


def normalize_ruleset(candidate: dict | None) -> tuple[dict, list[dict[str, str]]]:
    if not candidate:
        return copy.deepcopy(DEFAULT_RULESET), []
    return _merge_known(DEFAULT_RULESET, candidate)


def validate_ruleset(ruleset: dict) -> list[dict[str, str]]:
    issues = [
        _issue("/".join(map(str, err.path)), err.message)
        for err in sorted(_VALIDATOR.iter_errors(ruleset), key=lambda e: list(e.path))
    ]
    regex = ruleset.get("naming", {}).get("function_regex", "")
    if isinstance(regex, str):
        try:
            compiled = re.compile(regex)
            if not compiled.fullmatch("ltc2991_voltage_read"):
                issues.append(_issue("naming/function_regex", "regex should accept a module_object_action function"))
            if compiled.fullmatch("ltc2991_read"):
                issues.append(_issue("naming/function_regex", "regex should reject two-token function names"))
        except re.error as exc:
            issues.append(_issue("naming/function_regex", f"invalid regular expression: {exc}"))
    return issues


def _gate_errors(issues: list[dict[str, str]]) -> list[dict[str, str]]:
    return [issue for issue in issues if issue.get("severity", "error") == "error"]


def _sample_checks(ruleset: dict) -> list[dict[str, Any]]:
    checks: list[dict[str, Any]] = []
    regex = ruleset.get("naming", {}).get("function_regex", "")
    short_name_rejected = False
    regex_valid = True
    try:
        compiled = re.compile(regex)
        checks.append({
            "name": "function_regex_accepts_good_name",
            "passed": bool(compiled.fullmatch("ltc2991_voltage_read")),
            "detail": "ltc2991_voltage_read",
        })
        checks.append({
            "name": "function_regex_rejects_short_name",
            "passed": not bool(compiled.fullmatch("ltc2991_read")),
            "detail": "ltc2991_read",
        })
        short_name_rejected = not bool(compiled.fullmatch("ltc2991_read"))
    except re.error as exc:
        regex_valid = False
        checks.append({"name": "function_regex_compiles", "passed": False, "detail": str(exc)})
    if not regex_valid:
        return checks

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        drivers = root / "drivers"
        drivers.mkdir()
        good = drivers / "good.c"
        bad = drivers / "bad.c"
        good.write_text(
            '#include <stdio.h>\n\nint ltc2991_voltage_read(void)\n{\n    printf("ok\\r\\n");\n    return 0;\n}\n',
            encoding="utf-8",
        )
        bad.write_text(
            '#include <stdio.h>\n\nint ltc2991_read(void)\n{\n    printf("bad\\n");\n    return 0;\n}\n',
            encoding="utf-8",
        )
        good_violations = naming_linter.lint_file(good, ruleset, [drivers])
        bad_violations = naming_linter.lint_file(bad, ruleset, [drivers])
        good_errors = [v for v in good_violations if v.severity == "error"]
        bad_error_rules = {v.rule for v in bad_violations if v.severity == "error"}
        checks.append({
            "name": "sample_good_c_has_no_gate_errors",
            "passed": not good_errors,
            "detail": f"{len(good_errors)} error(s)",
        })
        checks.append({
            "name": "sample_bad_c_fails_print_rule",
            "passed": "print.line_terminator" in bad_error_rules,
            "detail": ", ".join(sorted(bad_error_rules)) or "no error rules",
        })
        checks.append({
            "name": "sample_bad_c_fails_function_rule",
            "passed": short_name_rejected,
            "detail": "regex rejects ltc2991_read",
        })
    return checks


def _flatten(data: dict, prefix: str = "") -> dict[str, Any]:
    out: dict[str, Any] = {}
    for key, value in data.items():
        path = f"{prefix}.{key}" if prefix else key
        if isinstance(value, dict):
            out.update(_flatten(value, path))
        else:
            out[path] = value
    return out


def diff_from_default(ruleset: dict) -> list[dict[str, Any]]:
    current = _flatten(ruleset)
    default = _flatten(DEFAULT_RULESET)
    diffs: list[dict[str, Any]] = []
    for path in sorted(set(current) | set(default)):
        old = default.get(path)
        new = current.get(path)
        if old != new:
            diffs.append({"path": path, "default": old, "candidate": new})
    return diffs


def evaluate_ruleset(candidate: dict | None) -> dict[str, Any]:
    normalized, normalize_issues = normalize_ruleset(candidate)
    issues = [*normalize_issues, *validate_ruleset(normalized)]
    checks = _sample_checks(normalized)
    for check in checks:
        if not check["passed"]:
            issues.append(_issue(f"checks/{check['name']}", check["detail"]))
    return {
        "ruleset": normalized,
        "valid": not _gate_errors(issues),
        "issues": issues,
        "diff": diff_from_default(normalized),
        "checks": checks,
    }


def build_candidate_from_text(source_text: str, llm: dict | None) -> dict[str, Any]:
    raw_candidate: dict | None = None
    llm_used = False
    llm_error = ""
    if llm and llm.get("enabled"):
        try:
            client = LlmClient(LlmConfig.resolve(llm))
            prompt = (
                "Convert this embedded-C coding standard into a Spec2Code ruleset JSON object. "
                "Return ONLY JSON. Use only these top-level keys: ruleset_version, formatting, naming, "
                "print_conventions, doxygen. Allowed formatting.brace_style values: allman, attach, k&r. "
                "Allowed formatting.indent: spaces_4. Allowed formatting.line_ending: crlf. "
                "Use defaults where the document is silent. Do not invent unsupported keys.\n\n"
                "DEFAULT RULESET:\n"
                f"{json.dumps(DEFAULT_RULESET, indent=2)}"
            )
            out = client.chat([
                {"role": "system", "content": prompt},
                {"role": "user", "content": source_text[:120_000]},
            ], temperature=0.0)
            raw_candidate = _json_object_from_text(out)
            llm_used = True
        except (LlmError, ValueError, json.JSONDecodeError) as exc:
            llm_error = str(exc)
    result = evaluate_ruleset(raw_candidate)
    result.update({
        "source_text": source_text,
        "llm_used": llm_used,
        "llm_error": llm_error,
        "needs_human_review": True,
    })
    if not llm_used and not llm_error:
        result["issues"] = [
            _issue("llm", "LLM is disabled; review and edit the default-derived ruleset manually", "warning"),
            *result["issues"],
        ]
    return result


def safe_ruleset_name(name: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9_.-]+", "_", name).strip("._")
    if not slug:
        slug = "company"
    if not slug.endswith(".ruleset.json"):
        slug += ".ruleset.json"
    return slug


def ruleset_write_path(filename: str) -> tuple[str, Path]:
    ref = f"std/user/{safe_ruleset_name(filename)}"
    override = os.environ.get("SPEC2CODE_RULESET_DIR")
    if override:
        path = Path(override) / safe_ruleset_name(filename)
    elif getattr(sys, "frozen", False):
        path = Path.cwd() / ref
    else:
        path = ROOT / ref
    return ref, path


def resolve_ruleset_ref(ref: str) -> Path | None:
    normalized = ref.replace("\\", "/").lstrip("/")
    path = Path(normalized)
    if path.is_absolute() or ".." in path.parts:
        return None
    candidates = [ROOT / normalized, Path.cwd() / normalized]
    if getattr(sys, "frozen", False):
        candidates.append(Path(sys.executable).resolve().parent / normalized)
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    return None


def save_ruleset(name: str, ruleset: dict) -> dict[str, Any]:
    result = evaluate_ruleset(ruleset)
    if not result["valid"]:
        return {"ok": False, **result}
    ref, path = ruleset_write_path(name)
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = json.dumps(result["ruleset"], indent=2) + "\n"
    path.write_text(payload, encoding="utf-8")
    return {"ok": True, "ref": ref, "path": str(path), **result}
